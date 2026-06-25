"""Client-ready export formats: CSV, Excel, PDF, Word (single-solution)."""
from __future__ import annotations

import csv
import io
import json
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.models.schemas import JobResponse

# Brand colors aligned with application theme
COLOR_PRIMARY = colors.HexColor("#000000")
COLOR_SECONDARY = colors.HexColor("#e7000b")
COLOR_TEXT = colors.HexColor("#ffffff")
COLOR_MUTED = colors.HexColor("#666666")

EXCEL_HEADER_FILL = PatternFill("solid", fgColor="000000")
EXCEL_HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
EXCEL_PASS_FILL = PatternFill("solid", fgColor="E8F5E9")
EXCEL_FAIL_FILL = PatternFill("solid", fgColor="FFEBEE")
EXCEL_NEUTRAL_FILL = PatternFill("solid", fgColor="FFF8E1")
EXCEL_TITLE_FONT = Font(bold=True, size=14, color="000000")
EXCEL_SECTION_FONT = Font(bold=True, size=12, color="E7000B")


def export_job_json(job: JobResponse) -> str:
    """Export single analysis result as JSON."""
    payload = job.model_dump(mode="json")
    # Clean up fields not relevant for single solution
    payload.pop("ranking", None)
    payload.pop("provider_groups", None)
    payload.pop("sarvam_batch_max_wait_seconds", None)
    return json.dumps(payload, indent=2, default=str)


def export_job_csv(job: JobResponse) -> str:
    """Export single analysis result as CSV."""
    output = io.StringIO()
    writer = csv.writer(output, lineterminator="\n")

    # Header section
    writer.writerow(["Analysis Report"])
    writer.writerow(["Job ID", job.job_id])
    writer.writerow(["Audio File", job.audio_filename])
    writer.writerow(["Call Reference", job.call_reference or "N/A"])
    writer.writerow(["Status", job.status.value])
    writer.writerow(["Timestamp", job.created_at or ""])
    writer.writerow([])

    # Analysis result
    if job.result:
        result = job.result
        writer.writerow(["Analysis Results"])
        writer.writerow(["Sentiment", result.sentiment])
        writer.writerow(["Confidence", result.confidence])
        writer.writerow(["Summary", result.summary])
        writer.writerow(["Recommended Action", result.recommended_action])
        writer.writerow([])
        
        if result.key_issues:
            writer.writerow(["Key Issues"])
            for issue in result.key_issues:
                writer.writerow([issue])
            writer.writerow([])
        
        if result.action_items:
            writer.writerow(["Action Items"])
            for action in result.action_items:
                writer.writerow([action])
            writer.writerow([])
        
        writer.writerow(["Transcript"])
        writer.writerow([result.transcript])
    else:
        writer.writerow(["Error", job.error or "No results available"])



def _auto_width_sheet(ws, min_width: int = 12, max_width: int = 48) -> None:
    for col_cells in ws.columns:
        length = min_width
        col_letter = get_column_letter(col_cells[0].column)
        for cell in col_cells:
            if cell.value:
                length = max(length, min(len(str(cell.value)) + 2, max_width))
        ws.column_dimensions[col_letter].width = length


def export_job_excel(job: JobResponse) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Analysis Report"

    # Title
    ws["A1"] = "Call Analysis Report"
    ws["A1"].font = EXCEL_TITLE_FONT
    ws.merge_cells("A1:D1")

    # Metadata
    row_idx = 3
    metadata = [
        ("Job ID", job.job_id),
        ("Audio File", job.audio_filename or "—"),
        ("Call Reference", job.call_reference or "—"),
        ("Status", job.status.value),
        ("Created", str(job.created_at or "—")),
        ("Completed", str(job.completed_at or "—")),
    ]
    for label, value in metadata:
        ws.cell(row=row_idx, column=1, value=label).font = Font(bold=True)
        ws.cell(row=row_idx, column=2, value=value)
        row_idx += 1

    row_idx += 1
    ws.cell(row=row_idx, column=1, value="Analysis Results").font = EXCEL_SECTION_FONT
    row_idx += 1

    if job.result:
        result = job.result
        analysis_data = [
            ("Sentiment", result.sentiment or "—"),
            ("Confidence", f"{result.confidence * 100:.0f}%" if result.confidence else "—"),
            ("Summary", result.summary or "—"),
            ("Issue Type", result.issue_type or "—"),
            ("Key Issues", "; ".join(result.key_issues) if result.key_issues else "—"),
            ("Action Items", "; ".join(result.action_items) if result.action_items else "—"),
            ("Escalation Risk", result.escalation_risk or "—"),
            ("Recommended Action", result.recommended_action or "—"),
        ]
        for label, value in analysis_data:
            ws.cell(row=row_idx, column=1, value=label).font = Font(bold=True)
            ws.cell(row=row_idx, column=2, value=value)
            row_idx += 1

        row_idx += 1
        ws.cell(row=row_idx, column=1, value="Transcript").font = EXCEL_SECTION_FONT
        row_idx += 1
        transcript_cell = ws.cell(row=row_idx, column=1, value=result.transcript or "—")
        transcript_cell.alignment = Alignment(wrap_text=True, vertical="top")
        ws.merge_cells(f"A{row_idx}:D{row_idx}")
    else:
        ws.cell(row=row_idx, column=1, value=job.error or "No results available")

    _auto_width_sheet(ws)
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def export_job_pdf(job: JobResponse) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=COLOR_PRIMARY,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle",
        parent=styles["Normal"],
        fontSize=11,
        textColor=COLOR_MUTED,
        spaceAfter=12,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=COLOR_SECONDARY,
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
    )

    story = []
    story.append(Paragraph("Call Analysis Report", title_style))
    report_date = job.completed_at or job.created_at or datetime.utcnow()
    date_str = report_date.strftime("%d %B %Y") if report_date else datetime.utcnow().strftime("%d %B %Y")
    story.append(Paragraph(f"Report Date: {date_str}", subtitle_style))
    story.append(Spacer(1, 0.15 * inch))

    # Job details
    story.append(Paragraph("Job Details", section_style))
    job_meta = [
        ["Job ID", job.job_id],
        ["Audio File", job.audio_filename or "—"],
        ["Call Reference", job.call_reference or "—"],
        ["Status", job.status.value],
        ["Created", str(job.created_at or "—")],
    ]
    job_table = Table(job_meta, colWidths=[2.0 * inch, 4.5 * inch])
    job_table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), COLOR_PRIMARY),
            ("TEXTCOLOR", (0, 0), (0, -1), COLOR_TEXT),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ])
    )
    story.append(job_table)
    story.append(Spacer(1, 0.25 * inch))

    if job.result:
        result = job.result
        story.append(Paragraph("Analysis Results", section_style))
        analysis_data = [
            ["Sentiment", result.sentiment or "—"],
            ["Confidence", f"{result.confidence * 100:.0f}%" if result.confidence else "—"],
            ["Summary", result.summary or "—"],
            ["Key Issues", "; ".join(result.key_issues) if result.key_issues else "—"],
            ["Escalation Risk", result.escalation_risk or "—"],
            ["Recommended Action", result.recommended_action or "—"],
        ]
        analysis_table = Table(analysis_data, colWidths=[2.0 * inch, 4.5 * inch])
        analysis_table.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), COLOR_PRIMARY),
                ("TEXTCOLOR", (0, 0), (0, -1), COLOR_TEXT),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
            ])
        )
        story.append(analysis_table)
        story.append(Spacer(1, 0.25 * inch))

        if result.transcript:
            story.append(Paragraph("Transcript", section_style))
            story.append(Paragraph(result.transcript[:2000], body_style))
    else:
        story.append(Paragraph(job.error or "No results available", body_style))

    doc.build(story)
    return buffer.getvalue()


def _add_word_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)

    doc.add_paragraph()


def export_job_word(job: JobResponse) -> bytes:
    doc = Document()

    title = doc.add_heading("Call Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    report_date = job.completed_at or job.created_at or datetime.utcnow()
    date_str = report_date.strftime("%d %B %Y") if report_date else datetime.utcnow().strftime("%d %B %Y")
    subtitle = doc.add_paragraph(f"Report Date: {date_str}")
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)

    doc.add_heading("Job Details", level=1)
    _add_word_table(
        doc,
        ["Field", "Value"],
        [
            ["Job ID", job.job_id],
            ["Audio File", job.audio_filename or "—"],
            ["Call Reference", job.call_reference or "—"],
            ["Status", job.status.value],
            ["Created", str(job.created_at or "—")],
        ],
    )

    if job.result:
        result = job.result
        doc.add_heading("Analysis Results", level=1)
        _add_word_table(
            doc,
            ["Field", "Value"],
            [
                ["Sentiment", result.sentiment or "—"],
                ["Confidence", f"{result.confidence * 100:.0f}%" if result.confidence else "—"],
                ["Summary", result.summary or "—"],
                ["Key Issues", "; ".join(result.key_issues) if result.key_issues else "—"],
                ["Action Items", "; ".join(result.action_items) if result.action_items else "—"],
                ["Escalation Risk", result.escalation_risk or "—"],
                ["Recommended Action", result.recommended_action or "—"],
            ],
        )

        if result.transcript:
            doc.add_heading("Transcript", level=1)
            doc.add_paragraph(result.transcript)
    else:
        doc.add_paragraph(job.error or "No results available")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
